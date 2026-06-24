#!/usr/bin/env python3
"""Generate a branded Raja Boot House quotation .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand palette
OXBLOOD = RGBColor(0x6B, 0x1F, 0x2A)
COGNAC = RGBColor(0x9C, 0x5B, 0x2E)
CHARCOAL = RGBColor(0x2B, 0x2B, 0x2B)
CREAM = "FBF6EF"
GREY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = CHARCOAL

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)

def set_row_bold_white(row):
    for c in row.cells:
        shade(c, "6B1F2A")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = WHITE

def heading(text, color=OXBLOOD, size=14, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def body(text, bold=False, color=CHARCOAL, size=10.5, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10.5)
    return p

# ---------- Header / Title block ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("RAJA BOOT HOUSE")
r.font.bold = True
r.font.size = Pt(26)
r.font.color.rgb = OXBLOOD

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Premium Leather Footwear  •  E-Commerce Platform")
rs.font.size = Pt(11)
rs.font.color.rgb = COGNAC
rs.italic = True

q = doc.add_paragraph()
q.alignment = WD_ALIGN_PARAGRAPH.CENTER
rq = q.add_run("QUOTATION")
rq.font.bold = True
rq.font.size = Pt(15)
rq.font.color.rgb = CHARCOAL
rq.font.underline = True

# Meta table
meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Prepared by", "Cosstech India"),
    ("Prepared for", "Raja Boot House"),
    ("Quotation No.", "RBH-ECOM-2026-001"),
    ("Date  /  Validity", "24 June 2026  /  15 days"),
]
for i, (k, v) in enumerate(meta_data):
    c0, c1 = meta.rows[i].cells
    c0.text = ""; c1.text = ""
    rk = c0.paragraphs[0].add_run(k); rk.font.bold = True; rk.font.color.rgb = OXBLOOD; rk.font.size = Pt(10)
    rv = c1.paragraphs[0].add_run(v); rv.font.size = Pt(10)
    shade(c0, CREAM)

doc.add_paragraph()

# ---------- 1. Summary ----------
heading("1.  Project Summary")
body("A complete, production-ready e-commerce platform for Raja Boot House — a premium "
     "leather footwear brand. Built on a modern, fast, SEO-optimised stack (Next.js App "
     "Router, MongoDB, Razorpay) and ready for deployment on Vercel with zero additional "
     "infrastructure cost.")
body("This quotation covers the full platform already delivered plus the essential "
     "completion items required to make it legally and operationally ready to go live — "
     "all delivered same-day.")

# ---------- 2. Delivered features ----------
heading("2.  Features Delivered (Already Built)")

sections = [
    ("A.  Storefront & Catalog", [
        "Product catalogue with size/colour variants, per-variant stock & images",
        "Filtering by category, brand, gender, occasion + full-text search",
        "Categories, brands, collections, banners, announcements",
        "Product detail pages with image gallery, share, recently-viewed",
        "SEO suite: server-rendered pages, JSON-LD, sitemap, robots rules",
        "Premium brand UI with smooth animations (mobile + desktop)",
    ]),
    ("B.  Cart, Checkout & Payments", [
        "Server-synced shopping cart",
        "3-step checkout (address - shipping - payment)",
        "Razorpay — UPI, Card, Net Banking, Wallet + Cash on Delivery",
        "Secure payment signature verification + webhook handling",
        "Server-side price re-computation (fraud-proof — client prices never trusted)",
        "Atomic stock deduction with rollback on failure",
        "Human-readable sequential order IDs (RBH-XXXXX)",
    ]),
    ("C.  Customer Accounts", [
        "Email/password login (encrypted) + Google sign-in",
        "Forgot/reset password via email",
        "Profile management, multiple saved addresses",
        "Order history + invoice view",
        "Wishlist",
        "Loyalty points — 5% earned on delivery, redeemable at checkout",
    ]),
    ("D.  Orders & Fulfilment", [
        "Full order lifecycle (Placed - Confirmed - Packed - Shipped - Out for Delivery - Delivered)",
        "Cancellation, return & refund states with status history & audit log",
        "Real Razorpay refunds to original payment method",
        "Customer self-cancel with automatic stock restoration",
    ]),
    ("E.  Marketing Tools", [
        "Coupons (flat / percentage / free-shipping, min-cart, usage limits, validity)",
        "Flash sales (time-limited, percentage or flat discounts)",
        "Product reviews & ratings (verified-purchase, helpful votes, photos)",
        "Newsletter subscribe/unsubscribe",
        "Automated abandoned-cart recovery emails",
    ]),
    ("F.  Admin Panel", [
        "Dashboard — live revenue, sales chart, top products, low-stock alerts",
        "Product / category / brand / collection management",
        "Bulk product import via CSV + CSV order export",
        "Order management, refund payouts export, printable shipping labels",
        "Coupons, flash sales, customers, reviews, CMS content, store settings",
        "Delivery-partner management + admin activity log",
    ]),
    ("G.  Vendor Portal", [
        "Vendor dashboard, products, orders, payouts & settings",
    ]),
    ("H.  Infrastructure", [
        "Image uploads via Vercel Blob (with local fallback)",
        "Transactional email (order confirmation, status updates, welcome)",
        "API rate-limiting, robust MongoDB Atlas connection handling",
        "Zero-config Vercel deployment",
    ]),
]
for title, items in sections:
    heading(title, color=COGNAC, size=11.5, space_before=8, space_after=3)
    for it in items:
        bullet(it)

# ---------- 3. Completion items ----------
heading("3.  Essential Completion Items (Delivered Today)")
body("Minimum gaps required for a safe, legal, production launch. All build quickly on the "
     "existing codebase and complete in the same delivery.", color=GREY, size=10)

comp = doc.add_table(rows=1, cols=3)
comp.style = "Table Grid"
hdr = comp.rows[0].cells
for i, h in enumerate(["#", "Item", "Why it matters"]):
    hdr[i].paragraphs[0].add_run(h)
set_row_bold_white(comp.rows[0])
comp_rows = [
    ("1", "Secrets hardening — remove live credentials from repo, rotate keys, add placeholder env template", "Security — prevents credential leak"),
    ("2", "GST persisted per order — freeze CGST/SGST + tax rate on each order & invoice at purchase time", "Legal — correct GST invoices & filing"),
    ("3", "Razorpay amount hardening — amount derived from saved order, never from client input", "Closes payment-tampering edge case"),
    ("4", "Push-notification cleanup — remove dead web-push stub & rotate exposed VAPID key", "Security + clean build"),
    ("5", "Order tracking link — surface delivery-partner tracking URL to customers", "Customer experience / less support load"),
    ("6", "Customer return-request flow — 'Request Return' action wired to existing return states", "Completes already-built returns backend"),
]
for r0 in comp_rows:
    cells = comp.add_row().cells
    for i, val in enumerate(r0):
        cells[i].paragraphs[0].add_run(val).font.size = Pt(9.5)
    cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ---------- 4. Commercials ----------
heading("4.  Commercials")
price = doc.add_table(rows=1, cols=2)
price.style = "Table Grid"
ph = price.rows[0].cells
ph[0].paragraphs[0].add_run("Description")
ph[1].paragraphs[0].add_run("Amount (Rs.)")
set_row_bold_white(price.rows[0])
price_rows = [
    ("Full e-commerce platform — all features in Section 2", "24,000"),
    ("Essential completion items — Section 3 (same-day)", "6,000"),
]
for desc, amt in price_rows:
    cells = price.add_row().cells
    cells[0].paragraphs[0].add_run(desc).font.size = Pt(10)
    rp = cells[1].paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run(amt).font.size = Pt(10)
# total row
trow = price.add_row().cells
trt = trow[0].paragraphs[0].add_run("TOTAL"); trt.font.bold = True; trt.font.color.rgb = WHITE
ta = trow[1].paragraphs[0]; ta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
tar = ta.add_run("Rs. 30,000"); tar.font.bold = True; tar.font.color.rgb = WHITE
for c in trow:
    shade(c, "9C5B2E")

body("Tax: As applicable (GST extra if registered).", size=9.5, color=GREY, after=2)
body("Payment terms: 50% advance, 50% on delivery / go-live.", size=9.5, color=GREY)

heading("Included", color=COGNAC, size=11, space_before=8, space_after=3)
for it in ["Same-day delivery of all items above", "Deployment to Vercel (production)", "Basic handover walkthrough"]:
    bullet(it)
heading("Not Included (available as add-ons)", color=COGNAC, size=11, space_before=6, space_after=3)
for it in [
    "SMS / WhatsApp transactional notifications (gateway charges extra)",
    "Courier API live-tracking integration (beyond static tracking link)",
    "Automated test suite",
    "Ongoing maintenance / AMC",
    "Third-party gateway, domain & hosting fees beyond Vercel free tier",
]:
    bullet(it)

# ---------- 5. Add-ons ----------
heading("5.  Optional Future Add-Ons (Separate Quote)")
for it in [
    "SMS/WhatsApp order & OTP notifications",
    "Live courier tracking + serviceability by pincode",
    "Back-in-stock & price-drop alerts",
    "Product Q&A, size guide, related-product recommendations",
    "Weight/zone-based shipping rate engine",
    "Analytics & conversion pixels (GA, Meta)",
    "Redis-backed rate limiting for scale",
    "Automated test coverage",
]:
    bullet(it)

# Footer note
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = foot.add_run("This quotation is an estimate based on current scope. Any change in scope "
                  "will be quoted separately.  —  Cosstech India")
rf.italic = True
rf.font.size = Pt(8.5)
rf.font.color.rgb = GREY

out = "/home/spideyo0/cosstechindia/raja-boot-house/Raja_Boot_House_Quotation.docx"
doc.save(out)
print("Saved:", out)
